from pathlib import Path

from docx import Document
from docx.shared import RGBColor


TEMPLATE = Path(r"C:\tmp\vkr_template.docx")
OUTPUT = Path(__file__).resolve().parent / "vkr_assignment_points_4_6_filled.docx"


POINT_4 = [
    "Анализ предметной области и существующих средств проектирования интерьера кухни",
    "Формирование требований к ИИ-агенту для автоматизированного планирования кухонного гарнитура",
    "Разработка системы правил и ограничений для размещения кухонных модулей и бытовой техники",
    "Проектирование архитектуры программного средства и алгоритмов генерации прямой и угловой кухни",
    "Реализация визуализации результата и проверка работы разработанного программного средства",
]

POINT_6 = [
    "Языки программирования: Python, TypeScript",
    "Фреймворки: FastAPI, React, Next.js",
    "Библиотеки и средства описания данных: Pydantic, стандартные библиотеки Python, SVG",
    "Форматы обмена и визуализации данных: JSON, SVG",
    "Инструменты сборки и запуска: npm, Uvicorn",
    "Сквозные технологии: Git, REST API, клиент-серверная архитектура, модульная архитектура",
]


def replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]

    for extra_run in paragraph.runs[1:]:
        extra_run.text = ""

    run.text = text
    run.font.color.rgb = RGBColor(0, 0, 0)


def main() -> None:
    doc = Document(TEMPLATE)

    for index, text in zip(range(26, 31), POINT_4):
        replace_paragraph_text(doc.paragraphs[index], text)

    for index, text in zip(range(35, 41), POINT_6):
        replace_paragraph_text(doc.paragraphs[index], text)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
