from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "kitchen_planner_unit_testing_report.docx"

SCREENSHOTS = {
    "backend_tests": ROOT / "Backend_успешный_запуск_тестов.png",
    "backend_terminal_coverage": ROOT / "Backend_покрытие_в_терминале.png",
    "backend_html_coverage": ROOT / "Backend_HTML-отчёт_покрытия.png",
    "frontend_tests": ROOT / "Frontend_успешный_запуск_тестов.png",
    "frontend_terminal_coverage": ROOT / "Frontend_покрытие_в_терминале.png",
    "frontend_html_coverage": ROOT / "Frontend_HTML-отчёт_покрытия.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=12, bold=False, italic=False, color=None, font="Times New Roman"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_centered(doc: Document, text="", size=12, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(1.25)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, italic=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "EAF1F8")
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=10, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_run_font(run, size=10)

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)

    doc.add_paragraph()
    return table


def add_code_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    for line_index, line in enumerate(lines):
        if line_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=10, font="Consolas")
    doc.add_paragraph()


def add_screenshot(doc: Document, key: str, caption: str, width_cm: float = 15.0):
    path = SCREENSHOTS[key]
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[name].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(13)


def add_title_page(doc: Document):
    add_centered(doc, "Министерство науки и высшего образования РФ")
    add_centered(doc, "Санкт-Петербургский политехнический университет Петра Великого")
    add_centered(doc, "Институт компьютерных наук и кибербезопасности")
    add_centered(doc, "Высшая школа программной инженерии")

    for _ in range(9):
        doc.add_paragraph()

    add_centered(doc, "ОТЧЁТ ПО КУРСОВОМУ ПРОЕКТУ", size=14, bold=True)
    doc.add_paragraph()
    add_centered(doc, " по дисциплине «Технологии разработки качественного программного обеспечения»")
    add_centered(doc, "Модульное тестирование веб-сайта Kitchen Planner")

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Выполнил")
    set_run_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("студент гр. 5130202/20202        Селиванов С.Ю.")
    set_run_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Руководитель     Маслаков А.П.")
    set_run_font(r)

    for _ in range(1):
        doc.add_paragraph()

    add_centered(doc, "Санкт-Петербург")
    add_centered(doc, "2026 г.")
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Описание выполненной работы", level=1)
    doc.add_heading("1.1. Общая характеристика проекта", level=2)
    add_body(
        doc,
        "Веб-приложение «Автоматизированный планировщик кухонного гарнитура» предназначено для "
        "формирования компоновки кухни на основании параметров помещения, выбранных пользователем "
        "модулей и набора конструктивных ограничений. Пользователь задаёт размеры кухни, тип "
        "планировки, бытовую технику, параметры верхних шкафов и дополнительные пожелания текстом. "
        "Система нормализует входные данные, фиксирует явно выбранные параметры, применяет правила "
        "автоматической оптимизации и возвращает набор модулей вместе с SVG-чертежами."
    )
    add_body(
        doc,
        "Архитектура проекта включает клиентскую часть на React/Next.js, API-маршруты Next.js, "
        "серверное приложение FastAPI, генератор прямой кухни, генератор угловой кухни, модуль "
        "оптимизации длин, подсистему формирования верхних и нижних модулей, обработчик текстовых "
        "пожеланий и SVG-рендерер. Отдельно реализована поддержка авторизации, личного кабинета и "
        "истории генераций пользователя."
    )
    add_body(
        doc,
        "Цель модульного тестирования на текущем этапе — проверить наиболее рискованные участки "
        "логики: разбор пользовательских пожеланий, группировку верхних шкафов, работу угловой "
        "генерации, формирование SVG-представлений и корректность проксирующих API-маршрутов frontend."
    )

    add_table(
        doc,
        ["Компонент", "Технологии", "Назначение"],
        [
            ["Frontend", "React, Next.js, TypeScript", "Интерфейс ввода параметров кухни, фиксация выбранных значений, вывод предупреждений, истории и SVG-эскизов."],
            ["API routes", "Next.js route handlers", "Проксирование запросов от интерфейса к backend-сервису генерации и распознавания пожеланий."],
            ["Backend", "Python, FastAPI", "HTTP API, генерация планировки, обработка истории, авторизация и интеграция с базой данных."],
            ["Генератор", "Python", "Построение прямой и угловой кухни, подбор модулей, оптимизация размеров и применение ограничений."],
            ["Рендерер", "Python, SVG", "Построение top view, front view и side view с размерами, фасадами, техникой и условными обозначениями."],
            ["База данных", "PostgreSQL, SQLAlchemy", "Хранение пользователей и истории генераций."],
            ["LLM-подсистема", "Ollama + правила", "Распознавание текстовых пожеланий пользователя и преобразование их в фиксируемые параметры формы."],
        ],
    )

    doc.add_heading("1.2. Используемые инструменты", level=2)
    add_table(
        doc,
        ["Инструмент", "Назначение"],
        [
            ["pytest", "Запуск модульных тестов backend-логики."],
            ["pytest-cov", "Формирование отчёта покрытия backend-кода."],
            ["Vitest", "Запуск модульных тестов frontend API-маршрутов."],
            ["@vitest/coverage-v8", "Формирование отчёта покрытия frontend-тестов."],
            ["jsdom", "Тестовая DOM-среда для frontend-тестирования."],
            ["React Testing Library", "Библиотека для будущего тестирования React-компонентов пользовательского интерфейса."],
            ["Next.js build", "Проверка компиляции и сборки frontend-приложения."],
            ["python compileall", "Проверка синтаксической корректности Python-модулей backend."],
            ["HTML coverage report", "Визуальная проверка покрытия тестами по файлам и строкам."],
        ],
    )

    doc.add_heading("1.3. Применённые техники тест-дизайна", level=2)
    doc.add_heading("1.3.1. Классы эквивалентности", level=3)
    add_body(
        doc,
        "Входные данные были разделены на классы эквивалентности по типам пользовательских параметров "
        "и ожидаемому поведению генератора. Такой подход позволяет проверять не каждую комбинацию "
        "параметров, а представителей классов, которые должны обрабатываться одинаково."
    )
    add_bullets(
        doc,
        [
            "тип кухни: прямая и угловая;",
            "тип верхних шкафов: распашные и подъёмные;",
            "размещение СВЧ: соло, встроенная в колонну, встроенная в верхний шкаф;",
            "верхние модули: обычный upper cabinet, сушка, шкаф вытяжки, антресоль;",
            "короткий верхний модуль: объединяемый с соседом и не объединяемый с соседом;",
            "текстовые пожелания: явно распознаваемые правилами, неявные фразы, некорректные значения;",
            "ответ backend: успешная генерация и ошибка engine.",
        ],
    )

    doc.add_heading("1.3.2. Граничные значения", level=3)
    add_body(
        doc,
        "Для модулей кухни важны граничные значения ширин: минимальные размеры drawer-модулей, "
        "минимальная ширина складного подъёмника 500 мм, ширины варочной поверхности 300/600/800/900 мм, "
        "ширины посудомоечной машины 450/600 мм. В тестах отдельно проверяется случай короткого "
        "верхнего сегмента 250 мм, который должен объединяться с соседним модулем либо переводиться "
        "в распашный фасад, если объединение невозможно."
    )

    doc.add_heading("1.3.3. Негативное тестирование", level=3)
    add_body(
        doc,
        "Негативные проверки направлены на ситуации, в которых система должна не падать, а корректно "
        "вернуть ошибку или безопасный fallback. Например, API-маршрут /api/plan должен обернуть "
        "ошибку backend engine в понятный JSON-ответ, а подсистема распознавания пожеланий должна "
        "игнорировать неизвестные поля и недопустимые значения."
    )

    doc.add_heading("1.3.4. Изоляция зависимостей", level=3)
    add_body(
        doc,
        "Frontend-тесты изолируют внешний backend с помощью mock-реализации fetch. Это позволяет "
        "проверять поведение API routes независимо от фактического запуска FastAPI-сервиса. Backend-тесты "
        "проверяют чистые функции генерации, рендера и нормализации параметров без обращения к базе данных "
        "и без запуска HTTP-сервера."
    )

    doc.add_heading("2. Отчёт о прохождении тестов и покрытии", level=1)
    doc.add_heading("2.1. Backend (Python / FastAPI)", level=2)
    add_body(
        doc,
        "Для backend-модуля создан набор unit-тестов в каталоге engine/tests. Проверяются правила "
        "группировки верхних модулей, fallback-поведение коротких подъёмников, работа генератора угловой "
        "кухни с верхней встроенной СВЧ, корректность формирования SVG-представлений и нормализация "
        "результатов распознавания пожеланий."
    )
    add_code_block(
        doc,
        [
            "cd C:\\Users\\Степан\\kitchen-planner\\engine",
            ".\\.venv\\Scripts\\python.exe -m pytest",
            ".\\.venv\\Scripts\\python.exe -m pytest --cov=app --cov-report=term --cov-report=html",
        ],
    )
    add_table(
        doc,
        ["Тестовый файл", "Проверяемая область", "Ключевые проверки"],
        [
            ["test_wall_modules.py", "Формирование верхних шкафов", "Объединение коротких lift-сегментов, объединение с сушкой, fallback в распашной фасад."],
            ["test_generator_renderer.py", "Генератор и рендерер", "Единая высота верхних шкафов при СВЧ в верхнем модуле, построение всех SVG-видов угловой кухни."],
            ["test_llm_preferences.py", "Распознавание пожеланий", "Извлечение JSON, фильтрация недопустимых patch-значений, определение подъёмных верхних фасадов из естественной фразы."],
        ],
    )
    add_body(
        doc,
        "По результатам запуска backend-набора выполнено 8 тестов, все тесты пройдены успешно. "
        "Общее покрытие backend-кода составило 63%. Покрытие распределено неравномерно: модули генерации "
        "и рендера покрыты лучше, чем авторизация и работа с базой данных, так как текущий отчёт посвящён "
        "в первую очередь логике планировщика."
    )
    add_screenshot(doc, "backend_tests", "Рисунок 1. Успешный запуск backend unit-тестов", width_cm=15.5)
    add_screenshot(doc, "backend_terminal_coverage", "Рисунок 2. Покрытие backend-кода в терминале", width_cm=15.5)
    add_screenshot(doc, "backend_html_coverage", "Рисунок 3. HTML-отчёт покрытия backend-кода", width_cm=11.0)

    doc.add_heading("2.2. Frontend (Next.js / TypeScript)", level=2)
    add_body(
        doc,
        "Для frontend-части создан набор unit-тестов API-маршрутов Next.js. На данном этапе проверяются "
        "маршруты /api/plan и /api/preferences/parse, так как они являются связующим слоем между "
        "пользовательским интерфейсом и backend engine. Внешний backend в тестах заменяется mock-реализацией fetch."
    )
    add_code_block(
        doc,
        [
            "cd C:\\Users\\Степан\\kitchen-planner\\web",
            "npm test",
            "npm run test:coverage",
        ],
    )
    add_table(
        doc,
        ["Тестовый файл", "Проверяемая область", "Ключевые проверки"],
        [
            ["api/plan/route.test.ts", "API-маршрут генерации", "Проксирование успешного ответа engine и корректная обработка неуспешного HTTP-ответа."],
            ["api/preferences/parse/route.test.ts", "API-маршрут распознавания пожеланий", "Передача текста пожеланий в backend parser и возврат patch-результата клиенту."],
        ],
    )
    add_body(
        doc,
        "По результатам запуска frontend-набора выполнено 3 теста, все тесты пройдены успешно. "
        "Покрытие протестированных API routes составило 87.5% по statements и lines, 100% по functions. "
        "Дальнейшее расширение набора должно включать тесты React-компонентов формы и личного кабинета."
    )
    add_screenshot(doc, "frontend_tests", "Рисунок 4. Успешный запуск frontend unit-тестов", width_cm=15.0)
    add_screenshot(doc, "frontend_terminal_coverage", "Рисунок 5. Покрытие frontend-кода в терминале", width_cm=14.0)
    add_screenshot(doc, "frontend_html_coverage", "Рисунок 6. HTML-отчёт покрытия frontend-кода", width_cm=15.5)

    doc.add_heading("2.3. Итоговая сводка покрытия", level=2)
    add_table(
        doc,
        ["Подсистема", "Команда", "Количество тестов", "Покрытие", "Комментарий"],
        [
            ["Backend", "pytest --cov=app", "8", "63%", "Покрыты ключевые правила генератора, wall_modules, renderer и LLM preferences."],
            ["Frontend", "vitest run --coverage", "3", "87.5%", "Покрыты API routes /api/plan и /api/preferences/parse."],
            ["Сборка frontend", "npm run build", "-", "успешно", "Next.js-приложение компилируется после добавления тестовой инфраструктуры."],
            ["Компиляция backend", "python -m compileall engine/app", "-", "успешно", "Python-модули проходят синтаксическую проверку."],
        ],
    )

    doc.add_heading("3. Процедура расширения тестового набора", level=1)
    doc.add_heading("3.1. Пример: расширение backend-тестов для нового правила верхних шкафов", level=2)
    add_body(
        doc,
        "При добавлении нового правила генерации рекомендуется сначала выделить чистую функцию или "
        "локальный сценарий, который можно проверить без запуска сервера. Например, для правила "
        "объединения коротких складных подъёмников был добавлен тест, передающий в apply_wall_facade_grouping "
        "короткий upper cabinet шириной 250 мм и соседний модуль шириной 600 мм."
    )
    add_code_block(
        doc,
        [
            "def test_lift_upper_merges_short_segment_with_regular_neighbor():",
            "    modules = [",
            "        _wall_module('upper_cabinet', 400, 250),",
            "        _wall_module('upper_cabinet', 650, 600),",
            "    ]",
            "    result = apply_wall_facade_grouping(modules, upper_cabinet_opening='lift')",
            "    assert result[0]['width_mm'] == 850",
            "    assert result[0]['facade_opening'] == 'lift_fold'",
        ],
    )
    add_body(
        doc,
        "Такой тест фиксирует ожидаемое поведение на уровне доменного правила и защищает его от "
        "регрессий при дальнейшей переработке угловой генерации или SVG-рендера."
    )

    doc.add_heading("3.2. Рекомендации по дальнейшему развитию тестов", level=2)
    add_bullets(
        doc,
        [
            "Добавить unit-тесты для length_optimizer: уменьшение мойки, варочной поверхности, посудомоечной машины и изменение вариаций духовки/СВЧ.",
            "Добавить тесты для lower_modules: cutlery drawer, filler drawer, drawer_1/drawer_2/drawer_3 и запрет нестандартного изменения базовых модулей.",
            "Добавить тесты для corner_generator: перенос модулей между сторонами, запрет соседства dishwasher/oven/upper microwave с corner base, выбор corner tall.",
            "Добавить тесты renderer: отсутствие подписей у скрытых модулей, корректные размерные стрелки, отсутствие наложений и корректная отрисовка фальш-планок.",
            "Добавить frontend component-тесты: изменение параметров формы, автоматическая фиксация выбранных пользователем модулей, история генераций, загрузка и удаление генерации.",
            "Добавить тесты авторизации и личного кабинета с mock-базой или тестовой PostgreSQL-базой.",
            "Добавить end-to-end тесты Playwright для сценариев: регистрация, генерация кухни, сохранение результата, открытие истории и скачивание PNG.",
        ],
    )

    doc.add_heading("Заключение", level=1)
    add_body(
        doc,
        "В результате выполненной работы для проекта подготовлена начальная инфраструктура модульного "
        "тестирования backend и frontend-частей. Тестовый набор покрывает наиболее критичные правила "
        "текущей реализации: группировку верхних модулей, поведение складных подъёмников, генерацию "
        "угловой кухни, SVG-рендеринг и API-прокси Next.js. Полученные HTML-отчёты покрытия могут "
        "использоваться для дальнейшего планирования расширения тестового набора."
    )
    add_body(
        doc,
        "Текущее покрытие backend составляет 63%, frontend API-слоя — 87.5%. Эти показатели достаточны "
        "для первичного отчёта и демонстрации применённых техник тестирования, однако дальнейшая работа "
        "должна быть направлена на покрытие авторизации, базы данных, UI-компонентов и end-to-end "
        "сценариев взаимодействия пользователя с системой."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
